import os
import json
from typing import Dict, List
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import matplotlib.patches as patches

# Simple version without OpenAI (for testing)
class MathQuestionGenerator:
    def __init__(self, use_openai=False, api_key=None):
        self.use_openai = use_openai
        if use_openai and api_key:
            import openai
            openai.api_key = api_key
            self.openai = openai
    
    def generate_questions_manually(self):
        """Generate questions without LLM (for testing)"""
        
        # Question 1: Combinations
        question1 = {
            "title": "School Lunch Combinations Assessment",
            "description": "Assessment on counting principles and combinations",
            "question": """Each student at Riverside Elementary School can choose a lunch combination 
consisting of 1 main dish, 1 side dish, and 1 drink. The table shows the available options:

Main Dish: Pizza, Burger, Chicken, Pasta
Side Dish: French Fries, Salad, Fruit Cup  
Drink: Milk, Water, Juice

How many different lunch combinations are possible?""",
            "instruction": "Select the correct answer from the options below.",
            "difficulty": "moderate",
            "order": 1,
            "options": ["Eleven", "Sixteen", "Twenty-four", "Thirty-six", "Forty-eight"],
            "correct_answer_index": 3,
            "explanation": "4 main dishes × 3 side dishes × 3 drinks = 36 combinations",
            "subject": "Quantitative Math",
            "unit": "Data Analysis & Probability",
            "topic": "Counting & Arrangement Problems",
            "plusmarks": 1
        }
        
        # Question 2: Geometry
        question2 = {
            "title": "Cylindrical Container Packing Assessment",
            "description": "Assessment on geometry and spatial reasoning",
            "question": """The side view of a rectangular box containing 8 tightly packed cylindrical cans 
is shown. Each can has a radius of 3 cm and height of 10 cm. The cans are arranged in 2 rows of 4 cans. 
Which dimensions are closest to those of the rectangular box?""",
            "instruction": "Select the correct answer from the options below.",
            "difficulty": "moderate",
            "order": 2,
            "options": [
                "$6 \\times 12 \\times 10$",
                "$10 \\times 12 \\times 24$",
                "$12 \\times 24 \\times 10$",
                "$18 \\times 24 \\times 10$",
                "$12 \\times 24 \\times 15$"
            ],
            "correct_answer_index": 2,
            "explanation": "Width: 2 rows × 6 cm diameter = 12 cm, Length: 4 cans × 6 cm = 24 cm, Height: 10 cm",
            "subject": "Quantitative Math",
            "unit": "Geometry and Measurement",
            "topic": "Solid Figures (Volume of Cubes)",
            "plusmarks": 1,
            "has_image": True
        }
        
        return [question1, question2]
    
    def generate_with_openai(self, base_question):
        """Generate using OpenAI API"""
        if not self.use_openai:
            return None
            
        prompt = f"""Generate a similar math question to: {base_question}
        Return as JSON with: title, description, question, instruction, options (5), 
        correct_answer_index, explanation, subject, unit, topic"""
        
        response = self.openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a math education expert."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7
        )
        
        return json.loads(response.choices[0].message.content)
    
    def create_cylinder_diagram(self, filename="cylinder_packing.png"):
        """Create a cylinder packing diagram"""
        fig, ax = plt.subplots(1, 1, figsize=(10, 6))
        
        # Draw rectangular box
        box = patches.Rectangle((0, 0), 24, 12, linewidth=2, 
                               edgecolor='black', facecolor='none')
        ax.add_patch(box)
        
        # Draw cylinders (top view - circles)
        radius = 3
        for row in range(2):
            for col in range(4):
                x = 3 + col * 6
                y = 3 + row * 6
                circle = patches.Circle((x, y), radius, linewidth=1,
                                       edgecolor='blue', facecolor='lightblue', alpha=0.7)
                ax.add_patch(circle)
                # Add center dot
                ax.plot(x, y, 'b.', markersize=8)
        
        # Add dimensions
        ax.annotate('', xy=(0, -1), xytext=(24, -1),
                   arrowprops=dict(arrowstyle='<->', color='red', lw=1.5))
        ax.text(12, -2, '24 cm', ha='center', color='red', fontsize=12)
        
        ax.annotate('', xy=(-1, 0), xytext=(-1, 12),
                   arrowprops=dict(arrowstyle='<->', color='red', lw=1.5))
        ax.text(-2.5, 6, '12 cm', ha='center', color='red', fontsize=12, rotation=90)
        
        ax.set_xlim(-3, 26)
        ax.set_ylim(-3, 14)
        ax.set_aspect('equal')
        ax.set_title('Top View: 8 Cylindrical Cans in Rectangular Box\n(Each can: radius = 3 cm, height = 10 cm)', 
                    fontsize=14)
        ax.grid(True, alpha=0.3)
        ax.set_xlabel('Length', fontsize=12)
        ax.set_ylabel('Width', fontsize=12)
        
        plt.tight_layout()
        plt.savefig(filename, dpi=150, bbox_inches='tight')
        plt.show()
        print(f"Diagram saved as {filename}")
        return filename
    
    def format_question(self, question):
        """Format question in the required output format"""
        output = f"""@title {question['title']}
@description {question['description']}

@question {question['question']}
@instruction {question['instruction']}
@difficulty {question['difficulty']}
@Order {question['order']}
"""
        
        for i, option in enumerate(question['options']):
            if i == question['correct_answer_index']:
                output += f"@@option {option}\n"
            else:
                output += f"@option {option}\n"
        
        output += f"""@explanation {question['explanation']}
@subject {question['subject']}
@unit {question['unit']}
@topic {question['topic']}
@plusmarks {question['plusmarks']}
"""
        return output
    
    def create_word_document(self, questions, filename="generated_math_questions.docx"):
        """Create Word document with questions"""
        doc = Document()
        doc.add_heading('Generated Math Questions', 0)
        
        for i, question in enumerate(questions, 1):
            doc.add_heading(f'Question {i}', level=1)
            
            # Add formatted question
            formatted = self.format_question(question)
            doc.add_paragraph(formatted)
            
            # Add image if needed
            if question.get('has_image'):
                image_file = self.create_cylinder_diagram(f"question_{i}_diagram.png")
                doc.add_paragraph('\nQuestion Diagram:')
                try:
                    doc.add_picture(image_file, width=Inches(5))
                except:
                    doc.add_paragraph('[Image: See generated diagram file]')
            
            if i < len(questions):
                doc.add_page_break()
        
        doc.save(filename)
        print(f"\n✅ Word document created: {filename}")
        return filename

# Main execution function
def main():
    print("=" * 60)
    print("MATH QUESTION GENERATOR")
    print("=" * 60)
    
    # Choose mode
    print("\nSelect mode:")
    print("1. Generate without OpenAI (demo mode)")
    print("2. Generate with OpenAI API")
    
    choice = input("\nEnter choice (1 or 2): ").strip()
    
    if choice == "2":
        api_key = input("Enter your OpenAI API key: ").strip()
        generator = MathQuestionGenerator(use_openai=True, api_key=api_key)
        print("\n✅ Using OpenAI API for generation")
    else:
        generator = MathQuestionGenerator(use_openai=False)
        print("\n✅ Using demo mode (no API needed)")
    
    # Generate questions
    print("\n📝 Generating questions...")
    questions = generator.generate_questions_manually()
    
    # Display questions
    print("\n" + "=" * 60)
    print("GENERATED QUESTIONS:")
    print("=" * 60)
    
    for i, q in enumerate(questions, 1):
        print(f"\n--- Question {i} ---")
        print(generator.format_question(q))
    
    # Create Word document
    print("\n📄 Creating Word document...")
    doc_file = generator.create_word_document(questions)
    
    # Save to text file as well
    with open("questions_output.txt", "w") as f:
        for i, q in enumerate(questions, 1):
            f.write(f"Question {i}\n")
            f.write("=" * 40 + "\n")
            f.write(generator.format_question(q))
            f.write("\n\n")
    print("✅ Text file created: questions_output.txt")
    
    print("\n" + "=" * 60)
    print("✅ GENERATION COMPLETE!")
    print("=" * 60)
    print("\nGenerated files:")
    print(f"1. {doc_file} - Word document")
    print("2. questions_output.txt - Text format")
    print("3. *.png files - Diagrams (if any)")

if __name__ == "__main__":
    main()